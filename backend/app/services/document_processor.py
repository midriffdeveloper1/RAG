import difflib
from pathlib import Path

import docx  # python-docx
from pypdf import PdfReader

from app.core.config import get_settings

settings = get_settings()


class UnsupportedFileTypeError(ValueError):
    pass


def extract_text(file_path: str, file_type: str) -> str:
    """Extract raw text from a PDF or DOCX file on disk."""
    if file_type == "pdf":
        return _extract_pdf_text(file_path)
    if file_type == "docx":
        return _extract_docx_text(file_path)
    if file_type == "doc":
        # TODO: legacy .doc isn't supported by python-docx. Either:
        #   (a) reject .doc at upload time (simplest), or
        #   (b) shell out to `libreoffice --headless --convert-to docx`
        #       before extraction.
        raise UnsupportedFileTypeError(
            "Legacy .doc files aren't supported yet — please upload .docx or .pdf."
        )
    raise UnsupportedFileTypeError(f"Unsupported file type: {file_type}")


def _extract_pdf_text(file_path: str) -> str:
    reader = PdfReader(file_path)
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages).strip()


def _extract_docx_text(file_path: str) -> str:
    document = docx.Document(file_path)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Tables often hold pricing/hours in real documents — include them too.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n\n".join(paragraphs).strip()


def chunk_text(
    text: str,
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
) -> list[str]:
    chunk_size = chunk_size or settings.chunk_size
    chunk_overlap = chunk_overlap or settings.chunk_overlap

    text = text.strip()
    if not text:
        return []
    min_boundary_fraction = 0.5

    chunks: list[str] = []
    start = 0
    text_length = len(text)

    while start < text_length:
        end = min(start + chunk_size, text_length)

        if end < text_length:
            boundary = max(
                text.rfind("\n\n", start, end),
                text.rfind(". ", start, end),
                text.rfind("\n", start, end),
            )
            min_boundary_pos = start + int(chunk_size * min_boundary_fraction)
            if boundary != -1 and boundary >= min_boundary_pos:
                end = boundary + 1

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        if end >= text_length:
            break

        # Guarantee real forward progress every iteration, regardless of
        # where the boundary search landed — this is the actual fix for
        # the near-duplicate cascade bug.
        min_next_start = start + max(1, chunk_size - chunk_overlap)
        start = max(end - chunk_overlap, min_next_start)

    return _deduplicate_chunks(chunks)


def _deduplicate_chunks(chunks: list[str], similarity_threshold: float = 0.9) -> list[str]:
    deduped: list[str] = []
    for chunk in chunks:
        is_duplicate = any(
            difflib.SequenceMatcher(None, chunk, prior).ratio() >= similarity_threshold
            for prior in deduped[-3:]
        )
        if not is_duplicate:
            deduped.append(chunk)
    return deduped


def infer_file_type(filename: str) -> str:
    suffix = Path(filename).suffix.lower().lstrip(".")
    return suffix